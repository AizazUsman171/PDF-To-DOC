import pdfplumber
from docx import Document
from docx.shared import Inches
from PIL import Image
import os
import logging

# Folder paths
UPLOAD_FOLDER = r'C:\Users\AMMAR COMPUTER\Downloads\new uzairs project\pdf_to_docx_updated\uploads'
OUTPUT_FOLDER = r'C:\Users\AMMAR COMPUTER\Downloads\new uzairs project\pdf_to_docx_updated\output'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def convert_pdf_to_docx_with_pdfplumber(pdf_path, output_docx_path):
    with pdfplumber.open(pdf_path) as pdf:
        doc = Document()

        for page_num, page in enumerate(pdf.pages):
            text = page.extract_text()
            images = page.images

            # Add text to the document
            if text:
                doc.add_paragraph(text)

            # Extract and add images to the document
            for img_index, img in enumerate(images):
                x0, y0, x1, y1 = img["x0"], img["y0"], img["x1"], img["y1"]
                img_obj = page.to_image()
                img_bytes = img_obj.original

                # Save the image temporarily
                image_path = f"temp_image_{page_num}_{img_index}.png"
                with open(image_path, "wb") as img_file:
                    img_bytes.save(img_file)

                # Add image to docx
                doc.add_picture(image_path, width=Inches(5))

                # Remove the temporary image file
                os.remove(image_path)

        # Save the final docx file
        doc.save(output_docx_path)
        logging.info(f"Conversion complete. The DOCX file is saved at {output_docx_path}")

def convert_pdfs_in_folder():
    for pdf_filename in os.listdir(UPLOAD_FOLDER):
        if pdf_filename.lower().endswith('.pdf'):
            pdf_path = os.path.join(UPLOAD_FOLDER, pdf_filename)
            base_filename = os.path.splitext(pdf_filename)[0]
            docx_path = os.path.join(OUTPUT_FOLDER, f"{base_filename}.docx")
            
            if os.path.exists(docx_path):
                logging.info(f"{pdf_filename} is already converted to {base_filename}.docx. Skipping.")
                continue
            
            try:
                logging.info(f"Converting {pdf_filename} to {base_filename}.docx...")
                convert_pdf_to_docx_with_pdfplumber(pdf_path, docx_path)
                logging.info(f"{pdf_filename} converted to {base_filename}.docx successfully.")
            except Exception as e:
                logging.error(f"Failed to convert {pdf_filename}: {str(e)}")

if __name__ == "__main__":
    convert_pdfs_in_folder()
